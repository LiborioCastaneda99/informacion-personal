from selenium import webdriver
from selenium.webdriver.common.by import By
import time
from docx import Document
from docx.shared import Inches

URL = "https://liboriocastaneda99.github.io/informacion-personal/"

options = webdriver.ChromeOptions()
options.add_experimental_option('excludeSwitches', ['enable-logging'])

# Casos de prueba (ID, Descripción, Pasos resumidos)
casos = {
    "CP01": {
        "descripcion": "Registro válido",
        "pasos": "Llenar todos los campos correctamente y enviar el formulario."
    },
    "CP02": {
        "descripcion": "Nombre vacío",
        "pasos": "Dejar el campo 'nombre' vacío y completar los demás campos."
    },
    "CP03": {
        "descripcion": "Correo inválido",
        "pasos": "Ingresar un correo con formato incorrecto (ejemplo: correo_invalido)."
    },
    "CP04": {
        "descripcion": "Documento con letras",
        "pasos": "Ingresar letras en el campo de número de documento."
    },
    "CP05": {
        "descripcion": "Teléfono vacío",
        "pasos": "Dejar el campo de teléfono vacío y completar los demás campos."
    },
    "CP06": {
        "descripcion": "Múltiples registros",
        "pasos": "Registrar tres usuarios diferentes en el sistema."
    },
}

# Guardamos resultados [(id, descripcion, pasos, resultado, screenshot)]
resultados = []


def setup_driver():
    driver = webdriver.Chrome(options=options)
    driver.get(URL)
    driver.maximize_window()
    return driver


def registrar_resultado(caso_id, exito, error=None, screenshot=None):
    desc = casos[caso_id]["descripcion"]
    pasos = casos[caso_id]["pasos"]
    if exito:
        resultado = f"✅ {caso_id} OK"
    else:
        resultado = f"❌ {caso_id} FAIL - {error}"
    resultados.append((caso_id, desc, pasos, resultado, screenshot))


# ===============================
# CASOS DE PRUEBA
# ===============================
def cp01_registro_valido():
    driver = setup_driver()
    try:
        driver.find_element(By.NAME, "nombre").send_keys("Carlos")
        driver.find_element(By.NAME, "apellidos").send_keys("Pérez Gómez")
        driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
        driver.find_element(By.NAME, "numeroDocumento").send_keys("123456789")
        driver.find_element(By.NAME, "fechaNacimiento").send_keys("1990-12-31")
        driver.find_element(By.NAME, "genero").send_keys("M")
        driver.find_element(By.NAME, "email").send_keys("carlos.perez@example.com")
        driver.find_element(By.NAME, "telefono").send_keys("3124567890")
        driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
        time.sleep(2)
        driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
        driver.find_element(By.NAME, "direccion").send_keys("Calle 10 #20-30")
        driver.find_element(By.NAME, "estadoCivil").send_keys("Soltero")
        driver.find_element(By.NAME, "ocupacion").send_keys("Ingeniero")
        driver.find_element(By.CSS_SELECTOR, "button.button").click()
        time.sleep(2)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) > 0

        screenshot = "CP01.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP01", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP01_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP01", False, e, screenshot)
    finally:
        driver.quit()


def cp02_nombre_vacio():
    driver = setup_driver()
    try:
        driver.find_element(By.NAME, "nombre").send_keys("")
        driver.find_element(By.NAME, "apellidos").send_keys("Pérez")
        driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
        driver.find_element(By.NAME, "numeroDocumento").send_keys("987654321")
        driver.find_element(By.NAME, "fechaNacimiento").send_keys("1990-01-01")
        driver.find_element(By.NAME, "genero").send_keys("M")
        driver.find_element(By.NAME, "email").send_keys("test@example.com")
        driver.find_element(By.NAME, "telefono").send_keys("3120000000")
        driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
        time.sleep(2)
        driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
        driver.find_element(By.NAME, "direccion").send_keys("Calle falsa 123")
        driver.find_element(By.NAME, "estadoCivil").send_keys("Soltero")
        driver.find_element(By.NAME, "ocupacion").send_keys("Tester")
        driver.find_element(By.CSS_SELECTOR, "button.button").click()
        time.sleep(2)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) == 0

        screenshot = "CP02.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP02", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP02_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP02", False, e, screenshot)
    finally:
        driver.quit()


def cp03_correo_invalido():
    driver = setup_driver()
    try:
        driver.find_element(By.NAME, "nombre").send_keys("Ana")
        driver.find_element(By.NAME, "apellidos").send_keys("Gómez")
        driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
        driver.find_element(By.NAME, "numeroDocumento").send_keys("111222333")
        driver.find_element(By.NAME, "fechaNacimiento").send_keys("1992-05-10")
        driver.find_element(By.NAME, "genero").send_keys("F")
        driver.find_element(By.NAME, "email").send_keys("correo_invalido")
        driver.find_element(By.NAME, "telefono").send_keys("3111111111")
        driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
        time.sleep(2)
        driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
        driver.find_element(By.NAME, "direccion").send_keys("Carrera 20")
        driver.find_element(By.NAME, "estadoCivil").send_keys("Casada")
        driver.find_element(By.NAME, "ocupacion").send_keys("Abogada")
        driver.find_element(By.CSS_SELECTOR, "button.button").click()
        time.sleep(2)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) == 0

        screenshot = "CP03.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP03", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP03_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP03", False, e, screenshot)
    finally:
        driver.quit()


def cp04_documento_letras():
    driver = setup_driver()
    try:
        driver.find_element(By.NAME, "nombre").send_keys("Luis")
        driver.find_element(By.NAME, "apellidos").send_keys("Ramírez")
        driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
        driver.find_element(By.NAME, "numeroDocumento").send_keys("ABC123XYZ")
        driver.find_element(By.NAME, "fechaNacimiento").send_keys("1991-07-15")
        driver.find_element(By.NAME, "genero").send_keys("M")
        driver.find_element(By.NAME, "email").send_keys("luis@example.com")
        driver.find_element(By.NAME, "telefono").send_keys("3205555555")
        driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
        time.sleep(2)
        driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
        driver.find_element(By.NAME, "direccion").send_keys("Cra 15 #22")
        driver.find_element(By.NAME, "estadoCivil").send_keys("Soltero")
        driver.find_element(By.NAME, "ocupacion").send_keys("Arquitecto")
        driver.find_element(By.CSS_SELECTOR, "button.button").click()
        time.sleep(2)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) == 0

        screenshot = "CP04.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP04", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP04_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP04", False, e, screenshot)
    finally:
        driver.quit()


def cp05_telefono_vacio():
    driver = setup_driver()
    try:
        driver.find_element(By.NAME, "nombre").send_keys("Marta")
        driver.find_element(By.NAME, "apellidos").send_keys("López")
        driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
        driver.find_element(By.NAME, "numeroDocumento").send_keys("444555666")
        driver.find_element(By.NAME, "fechaNacimiento").send_keys("1993-09-09")
        driver.find_element(By.NAME, "genero").send_keys("F")
        driver.find_element(By.NAME, "email").send_keys("marta@example.com")
        driver.find_element(By.NAME, "telefono").send_keys("")
        driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
        time.sleep(2)
        driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
        driver.find_element(By.NAME, "direccion").send_keys("Calle 45")
        driver.find_element(By.NAME, "estadoCivil").send_keys("Soltera")
        driver.find_element(By.NAME, "ocupacion").send_keys("Contadora")
        driver.find_element(By.CSS_SELECTOR, "button.button").click()
        time.sleep(2)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) == 0

        screenshot = "CP05.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP05", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP05_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP05", False, e, screenshot)
    finally:
        driver.quit()


def cp06_registro_multiple():
    driver = setup_driver()
    try:
        for i in range(3):
            driver.find_element(By.NAME, "nombre").send_keys(f"User{i}")
            driver.find_element(By.NAME, "apellidos").send_keys("Test")
            driver.find_element(By.NAME, "tipoDocumento").send_keys("CC")
            driver.find_element(By.NAME, "numeroDocumento").send_keys(f"1000{i}")
            driver.find_element(By.NAME, "fechaNacimiento").send_keys("1990-01-01")
            driver.find_element(By.NAME, "genero").send_keys("M")
            driver.find_element(By.NAME, "email").send_keys(f"user{i}@example.com")
            driver.find_element(By.NAME, "telefono").send_keys(f"31000000{i}")
            driver.find_element(By.NAME, "deptSel").send_keys("Antioquia")
            time.sleep(1)
            driver.find_element(By.NAME, "ciudad").send_keys("Medellín")
            driver.find_element(By.NAME, "direccion").send_keys(f"Calle {i}")
            driver.find_element(By.NAME, "estadoCivil").send_keys("Soltero")
            driver.find_element(By.NAME, "ocupacion").send_keys("Tester")
            driver.find_element(By.CSS_SELECTOR, "button.button").click()
            time.sleep(1)

        tarjetas = driver.find_elements(By.CLASS_NAME, "card")
        assert len(tarjetas) >= 3

        screenshot = "CP06.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP06", True, screenshot=screenshot)
    except Exception as e:
        screenshot = "CP06_error.png"
        driver.save_screenshot(screenshot)
        registrar_resultado("CP06", False, e, screenshot)
    finally:
        driver.quit()


# ===============================
# Generar reporte en Word
# ===============================
def generar_reporte_word():
    doc = Document()
    doc.add_heading("Reporte de Pruebas - Información Personal", level=1)
    doc.add_paragraph("Resultados de los casos de prueba ejecutados con Selenium:\n")

    for caso_id, desc, pasos, resultado, screenshot in resultados:
        doc.add_heading(caso_id, level=2)
        doc.add_paragraph(f"Descripción: {desc}")
        doc.add_paragraph(f"Pasos: {pasos}")
        doc.add_paragraph(f"Resultado: {resultado}")
        if screenshot:
            doc.add_picture(screenshot, width=Inches(4))
        doc.add_paragraph("\n")  # espacio

    doc.save("reporte_pruebas.docx")
    print("\n📄 Reporte guardado en 'reporte_pruebas.docx'")


# ===============================
# Ejecutar todos los casos
# ===============================
if __name__ == "__main__":
    cp01_registro_valido()
    cp02_nombre_vacio()
    cp03_correo_invalido()
    cp04_documento_letras()
    cp05_telefono_vacio()
    cp06_registro_multiple()

    print("\n===== RESULTADOS =====")
    for r in resultados:
        print(r[3])  # solo el resultado

    generar_reporte_word()

    print("\n🎉 Pruebas completadas con exito!")